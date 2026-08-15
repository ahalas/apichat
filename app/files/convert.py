"""Convert Office attachments to text/CSV when a message is sent."""

from __future__ import annotations

import base64
import csv
import io
import re
from typing import Any

from app.providers.messages import MAX_ATTACHMENTS


class ConversionError(ValueError):
    """Raised when a Word or Excel file cannot be converted."""


def _ext(filename: str) -> str:
    if "." not in (filename or ""):
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


def _stem(filename: str) -> str:
    name = filename or "file"
    if "." in name:
        return name.rsplit(".", 1)[0]
    return name


def _safe_sheet_name(name: str) -> str:
    cleaned = re.sub(r"[^\w.-]+", "_", name or "sheet").strip("._") or "sheet"
    return cleaned[:40]


def office_kind(filename: str, mime: str = "") -> str | None:
    ext = _ext(filename)
    mime = (mime or "").lower()
    if ext == ".docx" or mime.endswith("wordprocessingml.document"):
        return "docx"
    if ext == ".xlsx" or mime.endswith("spreadsheetml.sheet"):
        return "xlsx"
    if ext == ".xls" or mime == "application/vnd.ms-excel":
        return "xls"
    return None


def convert_office_attachments(attachments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Replace .docx/.xlsx/.xls with .txt/.csv. Other attachments pass through."""
    out: list[dict[str, Any]] = []
    for att in attachments:
        kind = office_kind(str(att.get("filename") or ""), str(att.get("mime") or ""))
        try:
            data = base64.b64decode(str(att.get("data_base64") or ""), validate=False)
        except Exception as exc:
            raise ConversionError(f"Could not read {att.get('filename') or 'file'}.") from exc
        if kind == "docx":
            _append_within_cap(out, [_docx_to_txt(str(att.get("filename") or "document.docx"), data)])
        elif kind in {"xlsx", "xls"}:
            sheets = _spreadsheet_to_csvs(str(att.get("filename") or "workbook.xlsx"), data, kind)
            _append_sheets(out, sheets)
        else:
            _append_within_cap(out, [att])
    return out


def _append_within_cap(out: list[dict[str, Any]], items: list[dict[str, Any]]) -> None:
    if len(out) + len(items) > MAX_ATTACHMENTS:
        raise ConversionError(f"Too many attachments (max {MAX_ATTACHMENTS}).")
    out.extend(items)


def _append_sheets(out: list[dict[str, Any]], sheets: list[dict[str, Any]]) -> None:
    remaining = MAX_ATTACHMENTS - len(out)
    if remaining <= 0:
        raise ConversionError(f"Too many attachments (max {MAX_ATTACHMENTS}).")
    if len(sheets) <= remaining:
        out.extend(sheets)
        return
    if remaining == 1:
        out.append(_merge_csv_attachments(sheets))
        return
    out.extend(sheets[: remaining - 1])
    out.append(_merge_csv_attachments(sheets[remaining - 1 :]))


def _docx_to_txt(filename: str, data: bytes) -> dict[str, Any]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ConversionError("python-docx is required to convert Word files.") from exc
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise ConversionError(f"Could not convert {filename} to text.") from exc
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").rstrip()
        if text:
            parts.append(text)
    for table in doc.tables:
        if parts and parts[-1] != "":
            parts.append("")
        for row in table.rows:
            cells = [(cell.text or "").replace("\n", " ").strip() for cell in row.cells]
            parts.append("\t".join(cells))
        parts.append("")
    body = "\n".join(parts).strip()
    encoded = base64.b64encode(body.encode("utf-8")).decode("ascii")
    return {
        "filename": f"{_stem(filename)}.txt",
        "mime": "text/plain",
        "data_base64": encoded,
    }


def _spreadsheet_to_csvs(filename: str, data: bytes, kind: str) -> list[dict[str, Any]]:
    try:
        sheets = _xlsx_sheets(data) if kind == "xlsx" else _xls_sheets(data)
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"Could not convert {filename} to CSV.") from exc
    if not sheets:
        raise ConversionError(f"{filename} has no sheets to convert.")
    stem = _stem(filename)
    out: list[dict[str, Any]] = []
    for name, rows in sheets:
        csv_text = _rows_to_csv(rows)
        encoded = base64.b64encode(csv_text.encode("utf-8")).decode("ascii")
        label = _safe_sheet_name(name)
        out.append(
            {
                "filename": f"{stem}_{label}.csv" if len(sheets) > 1 else f"{stem}.csv",
                "mime": "text/csv",
                "data_base64": encoded,
            }
        )
    return out


def _xlsx_sheets(data: bytes) -> list[tuple[str, list[list[str]]]]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ConversionError("openpyxl is required to convert Excel files.") from exc
    workbook = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    try:
        sheets: list[tuple[str, list[list[str]]]] = []
        for worksheet in workbook.worksheets:
            rows = [[_cell_text(value) for value in row] for row in worksheet.iter_rows(values_only=True)]
            sheets.append((worksheet.title or "sheet", _trim_empty_rows(rows)))
        return sheets
    finally:
        workbook.close()


def _xls_sheets(data: bytes) -> list[tuple[str, list[list[str]]]]:
    try:
        import xlrd
    except ImportError as exc:
        raise ConversionError("xlrd is required to convert .xls files.") from exc
    book = xlrd.open_workbook(file_contents=data)
    sheets: list[tuple[str, list[list[str]]]] = []
    for worksheet in book.sheets():
        rows = [
            [_cell_text(worksheet.cell_value(r, c)) for c in range(worksheet.ncols)]
            for r in range(worksheet.nrows)
        ]
        sheets.append((worksheet.name or "sheet", _trim_empty_rows(rows)))
    return sheets


def _cell_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def _trim_empty_rows(rows: list[list[str]]) -> list[list[str]]:
    while rows and not any(cell.strip() for cell in rows[-1]):
        rows.pop()
    return rows


def _rows_to_csv(rows: list[list[str]]) -> str:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerows(rows)
    return buf.getvalue()


def _merge_csv_attachments(items: list[dict[str, Any]]) -> dict[str, Any]:
    chunks: list[str] = []
    for item in items:
        name = str(item.get("filename") or "sheet.csv")
        try:
            text = base64.b64decode(str(item.get("data_base64") or ""), validate=False).decode("utf-8")
        except Exception:
            text = ""
        chunks.append(f"# {name}\n{text.rstrip()}\n")
    first = items[0] if items else {"filename": "workbook.csv"}
    merged = "\n".join(chunks).strip() + "\n"
    return {
        "filename": str(first.get("filename") or "workbook.csv"),
        "mime": "text/csv",
        "data_base64": base64.b64encode(merged.encode("utf-8")).decode("ascii"),
    }
